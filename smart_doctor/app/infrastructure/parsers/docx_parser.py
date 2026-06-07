"""
DOCX 文档解析器（v2.1）

使用 python-docx 提取文本，支持：
- 段落文本提取
- 表格内容提取
- 标题层级识别
"""
from __future__ import annotations

import logging
import os
import time

from app.infrastructure.parsers.base import DocumentParser, ParsedDocument, TextSegment

logger = logging.getLogger(__name__)


class DocxParser(DocumentParser):
    supported_types = {"docx", "doc"}

    async def parse(self, file_path: str, **kwargs) -> ParsedDocument:
        t0 = time.time()
        file_size = _safe_size(file_path)

        # .doc 旧格式不支持，降级为仅文件名信息
        if file_path.lower().endswith(".doc") and not file_path.lower().endswith(".docx"):
            return _error_doc(
                "旧版 .doc 格式暂不支持，请转为 .docx 格式后再上传", file_size, t0
            )

        try:
            from docx import Document
        except ImportError:
            return _error_doc(
                "python-docx 未安装，请执行 pip install python-docx", file_size, t0
            )

        segments: list[TextSegment] = []
        page_estimate = 0  # DOCX 不直接暴露页码，按段落数估算

        try:
            doc = Document(file_path)

            current_heading = None
            para_count = 0
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                # 检测标题
                if para.style.name.startswith("Heading"):
                    current_heading = text
                elif "heading" in para.style.name.lower():
                    current_heading = text

                page_estimate = (para_count // 40) + 1  # 约 40 段落≈1 页
                segments.append(TextSegment(
                    text=text,
                    page=page_estimate,
                    heading=current_heading,
                    confidence=0.95,
                ))
                para_count += 1

            # 提取表格
            for ti, table in enumerate(doc.tables):
                for ri, row in enumerate(table.rows):
                    cells = [cell.text.strip() for cell in row.cells]
                    row_text = " | ".join(c for c in cells if c)
                    if row_text:
                        segments.append(TextSegment(
                            text=row_text,
                            page=page_estimate + ti + 1,
                            row=ri + 1,
                            heading=f"表格{ti + 1}",
                            confidence=0.90,
                        ))
        except Exception as e:
            return _error_doc(str(e), file_size, t0)

        elapsed = (time.time() - t0) * 1000
        logger.info(
            "DOCX parsed: path=%s paragraphs=%d segments=%d time=%.0fms",
            file_path, para_count, len(segments), elapsed,
        )

        full_text = "\n".join(s.text for s in segments)
        return ParsedDocument(
            text=full_text,
            segments=segments,
            page_count=page_estimate,
            encoding="utf-8",
            parse_method="python-docx",
            parse_duration_ms=elapsed,
            file_size=file_size,
            file_type="docx",
        )


def _safe_size(path: str) -> int:
    try:
        return os.path.getsize(path)
    except OSError:
        return 0


def _error_doc(msg: str, file_size: int, t0: float) -> ParsedDocument:
    return ParsedDocument(
        text="",
        parse_method="python-docx",
        parse_duration_ms=(time.time() - t0) * 1000,
        file_size=file_size,
        file_type="docx",
        error=msg,
    )
